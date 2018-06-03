import docx
import pandas as pd
doc = docx.Document('SampleInputDoc3-Hardware Problems.docx')
questions=[]
ans=[]
for i in range(len(doc.paragraphs)):
    if (doc.paragraphs[i].text=="Symptom" and i<(len(doc.paragraphs)-1) and doc.paragraphs[i+1].text!=""):
        questions.append(doc.paragraphs[i+1].text)

for i in range(len(doc.paragraphs)):
    if (doc.paragraphs[i].text == "Symptom" and i < (len(doc.paragraphs) - 2)):
        j=i+2
        new_ans = ""
        while(doc.paragraphs[j].text!="Symptom" and j < (len(doc.paragraphs) - 1)):
            new_ans=new_ans+'\n*'+doc.paragraphs[j].text
            j+=1
        ans.append(new_ans)

help=[]
app=[]
for i in range(len(questions)):
    help.append('genericTroubleshoot')
    app.append(0)

data = {'Title':questions,'Resolution':ans,'HelpTopic':help,'approved':app}
df = pd.DataFrame(data,columns=data.keys())
print(df)
df.to_csv('inp3.csv')
