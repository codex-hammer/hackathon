from flask import request
from flask import Flask,render_template, jsonify,json
import warnings
warnings.filterwarnings(action='ignore', category=UserWarning, module='gensim')
import gensim
from gensim.test.utils import get_tmpfile
import string
from nltk.tokenize import word_tokenize
import pandas as pd
from nltk.tokenize import RegexpTokenizer
import numpy as np
import nltk
from nltk.stem import WordNetLemmatizer
wordnet_lemmatizer = WordNetLemmatizer()
from nltk.corpus import stopwords
stop_words = set(stopwords.words('english'))
tokenizer = RegexpTokenizer(r'\w+')

import docx
doc = docx.Document('NecessaryDox/SampleInputDoc3-Hardware Problems.docx')

df=pd.read_csv('NecessaryDox/finalInput.csv',encoding = "ISO-8859-1")

titles = df['Title']
answers = df['Resolution']

###########################################################################################################

questions=[]
ans=[]
for i in range(len(doc.paragraphs)):
    if (doc.paragraphs[i].text=="Symptom" and i<(len(doc.paragraphs)-1) and doc.paragraphs[i+1].text!=""):
        questions.append(doc.paragraphs[i+1].text)

for i in range(len(doc.paragraphs)):
    if (doc.paragraphs[i].text == "Symptom" and i < (len(doc.paragraphs) - 2)):
        j=i+2
        new_ans = ""
        while(doc.paragraphs[j].text!="Symptom" and j < (len(doc.paragraphs) - 1)):
            new_ans=new_ans+'\n*'+doc.paragraphs[j].text
            j+=1
        ans.append(new_ans)

help=[]
app=[]
for i in range(len(questions)):
    help.append('genericTroubleshoot')
    app.append(0)



app = Flask(__name__)

@app.route('/')
def main():
    return render_template('main.html')

@app.route('/aaa')
def showmain():
    return render_template('index.html')

@app.route('/aaa2')
def showmain2():
    return render_template('index2.html')

@app.route('/ans', methods=['POST'])
def give_answer():
    gen_docs = [[w.lower() for w in tokenizer.tokenize(text)] for text in titles]
    titles_train2 = [[x for x in item if not x in stop_words] for item in gen_docs]
    titles_train = [[wordnet_lemmatizer.lemmatize(x) for x in item] for item in titles_train2]
    dictionary1 = gensim.corpora.Dictionary(titles_train)
    dictionary1.save_as_text('buffers/hackdictionary.txt')
    corpus = [dictionary1.doc2bow(gen_doc) for gen_doc in titles_train]
    gensim.corpora.MmCorpus.serialize('buffers/corphack.mm', corpus)

    tf_idf = gensim.models.TfidfModel(corpus)

    sims = gensim.similarities.Similarity('buffers/simshack', tf_idf[corpus], num_features=len(dictionary1))
    ques=dict(title=request.form['name'])
    query_doc = [w.lower() for w in tokenizer.tokenize(ques['title'])]
    query_doc_bow = dictionary1.doc2bow(query_doc)
    query_doc_tf_idf = tf_idf[query_doc_bow]
    x = sims[query_doc_tf_idf]
    x = np.array(x)
    print(x)
    indices = x.argsort()[-10:][::-1]
    print(indices)
    titles_in_consideration = [titles[i] for i in indices]
    probable_answers = [answers[i] for i in indices]
    print(titles_in_consideration)
    print('-----------------------------------------------------------------------------------')
    #################################################################################################

    corpus2 = titles_in_consideration
    gen_docs1 = [[w.lower() for w in tokenizer.tokenize(text)] for text in corpus2]
    titles_train1 = [[wordnet_lemmatizer.lemmatize(x) for x in item] for item in gen_docs1]
    dictionary2 = gensim.corpora.Dictionary(titles_train1)
    corpus1 = [dictionary2.doc2bow(gen_doc) for gen_doc in titles_train1]
    lsi = gensim.models.LsiModel(corpus1, id2word=dictionary2)
    sims1 = gensim.similarities.MatrixSimilarity(lsi[corpus1])
    query1 = ques['title']
    query_doc1 = [w.lower() for w in tokenizer.tokenize(query1)]
    query_doc_bow1 = dictionary2.doc2bow(query_doc1)
    query_doc_tf_idf1 = lsi[query_doc_bow1]
    x = sims1[query_doc_tf_idf1]
    x = np.array(x)
    final2 = x.argmax()
    print(titles_in_consideration[final2])
    final_answer = probable_answers[final2]
    print(final_answer)
    msg = ""
    for i in final_answer:
        if i=='*':
            msg=msg+'<br>*'
        else:
            msg=msg+i
    #--------------------------------------------------------------------
    f = open("templates/answer.html",'w',encoding='utf-8')
    message = "<html><body bgcolor='#d089f9'><center><font face='Sans-serif'><h2>Hello! Hope this helps</center></h2><br><br><br><div style='width:1300px; border:4px solid white; padding: 15px 15px 15px 15px;'><h2>"+msg+"</h2></div></font><br><br><center><a href='/'><h1>Go Home</h1></a><a href='aaa2'><h1>Not helping? Try Troubleshooting.</h1></a></center></body></html>"
    f.write(message)
    f.close()
    #--------------------------------------------------------------------
    return final_answer
    #return json.dumps({'html': '<span>'+final_answer+'</span>'})

@app.route('/answer')
def showresult():
    return render_template('answer.html')

@app.route('/aaaaaa')
def showresult2():
    return render_template('answer2.html')

@app.route('/ans2', methods=['POST'])
def give_answer2():
    gen_docs = [[w.lower() for w in tokenizer.tokenize(text)] for text in questions]
    titles_train2 = [[x for x in item if not x in stop_words] for item in gen_docs]
    titles_train = [[wordnet_lemmatizer.lemmatize(x) for x in item] for item in titles_train2]
    dictionary1 = gensim.corpora.Dictionary(titles_train)
    dictionary1.save_as_text('buffers/hackdictionary2.txt')
    corpus = [dictionary1.doc2bow(gen_doc) for gen_doc in titles_train]
    gensim.corpora.MmCorpus.serialize('buffers/corphack2.mm', corpus)

    tf_idf = gensim.models.TfidfModel(corpus)

    sims = gensim.similarities.Similarity('buffers/simshack2', tf_idf[corpus], num_features=len(dictionary1))
    ques2=dict(title=request.form['fname'])
    query_doc = [w.lower() for w in tokenizer.tokenize(ques2['title'])]
    query_doc_bow = dictionary1.doc2bow(query_doc)
    query_doc_tf_idf = tf_idf[query_doc_bow]
    x = sims[query_doc_tf_idf]
    x = np.array(x)
    print(x)
    indices = x.argsort()[-10:][::-1]
    print(indices)
    titles_in_consideration = [questions[i] for i in indices]
    prob_answers = [ans[i] for i in indices]
    print(titles_in_consideration)
    print('-----------------------------------------------------------------------------------')
    #################################################################################################

    corpus2 = titles_in_consideration
    gen_docs1 = [[w.lower() for w in tokenizer.tokenize(text)] for text in corpus2]
    titles_train1 = [[wordnet_lemmatizer.lemmatize(x) for x in item] for item in gen_docs1]
    dictionary2 = gensim.corpora.Dictionary(titles_train1)
    corpus1 = [dictionary2.doc2bow(gen_doc) for gen_doc in titles_train1]
    lsi = gensim.models.LsiModel(corpus1, id2word=dictionary2)
    sims1 = gensim.similarities.MatrixSimilarity(lsi[corpus1])
    query1 = ques2['title']
    query_doc1 = [w.lower() for w in tokenizer.tokenize(query1)]
    query_doc_bow1 = dictionary2.doc2bow(query_doc1)
    query_doc_tf_idf1 = lsi[query_doc_bow1]
    x = sims1[query_doc_tf_idf1]
    x = np.array(x)
    final2 = x.argmax()
    print(titles_in_consideration[final2])
    final_answer = prob_answers[final2]
    print(final_answer)
    msg = ""
    for i in final_answer:
        if i=='*':
            msg=msg+'<br>*'
        else:
            msg=msg+i
    #--------------------------------------------------------------------
    f = open("templates/answer2.html",'w',encoding='utf-8')
    message = "<html><body bgcolor='#d089f9'><center><h2><font face='Sans-serif'>You can follow these steps to troubleshoot your problem.</center></h2><br><br><br><div style='width:1300px; border:4px solid white; padding: 15px 15px 15px 15px;'><h2>"+msg+"</h2></div></font><br><br><center><a href='/'><h1>Go Home</h1></a><a href='aaa'><h1>Not helping? Try Querying.</h1></a></center></body></html>"
    f.write(message)
    f.close()
    #--------------------------------------------------------------------
    return final_answer
    #return json.dumps({'html': '<span>'+final_answer+'</span>'})

if __name__ == '__main__':
    app.run(port=80, debug=True)

